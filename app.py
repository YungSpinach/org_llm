import os
import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from transformers import AutoModelForCausalLM, AutoTokenizer
from docx import Document as DocxDocument  # For Word documents
from langchain.schema import Document  # LangChain Document class
import torch

# Load documents
def load_documents(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            documents.append(Document(page_content=text))  # Wrap text in a Document object
        elif filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as file:
                text = file.read()
            documents.append(Document(page_content=text))  # Wrap text in a Document object
        elif filename.endswith(".docx"):
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])  # Extract text from Word document
            documents.append(Document(page_content=text))  # Wrap text in a Document object
    return documents

# Process documents
def process_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_documents(documents)  # Split into chunks
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    embeddings = embedding_model.encode([chunk.page_content for chunk in chunks])  # Generate embeddings
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)  # Create FAISS index
    index.add(np.array(embeddings))  # Add embeddings to index
    return chunks, index, embedding_model

# Load LLM
def load_llm(model_name="facebook/opt-1.3b"):
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForCausalLM.from_pretrained(model_name)
    return tokenizer, model

# def load_llm(model_name="deepseek-ai/DeepSeek-R1"):
#     tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
#     model = AutoModelForCausalLM.from_pretrained(
#         model_name,
#         trust_remote_code=True,
#         torch_dtype=torch.float16,  # Use FP16 for better performance
#         device_map="auto"  # Automatically map the model to available devices (CPU/GPU)
#     )
#     return tokenizer, model

# Semantic search
def search(query, index, chunks, embedding_model, top_k=3):
    query_embedding = embedding_model.encode([query])  # Encode query
    distances, indices = index.search(np.array(query_embedding), top_k)  # Search FAISS index
    return [chunks[i].page_content for i in indices[0]]  # Return top-k relevant chunks

# Generate response
def generate_response(query, context, tokenizer, model):
    input_text = f"Context: {context}\n\nQuestion: {query}\n\nAnswer:"
    inputs = tokenizer(input_text, return_tensors="pt")  # Tokenize input
    if torch.cuda.is_available():
        inputs = {key: value.to("cuda") for key, value in inputs.items()}  # Move inputs to GPU if available
    outputs = model.generate(**inputs, max_length=500)  # Generate response
    return tokenizer.decode(outputs[0], skip_special_tokens=True)  # Decode response

# Streamlit App
st.title("LLM Document Querying App")

# Preload documents
folder_path = "C:\\Users\\Christian.Jaccarini\\OneDrive - NEW ECONOMICS FOUNDATION\\2. Projects\\NEF_knowledge_base\\org_llm\\documents"  # Replace with your folder path
documents = load_documents(folder_path)
if not documents:
    st.error("No documents found in the specified folder. Please check the folder path.")
    st.stop()

# Process documents
chunks, index, embedding_model = process_documents(documents)

# Load LLM
tokenizer, model = load_llm()

# Query interface
query = st.text_input("Enter your query:")
if query:
    try:
        relevant_chunks = search(query, index, chunks, embedding_model)  # Find relevant chunks
        context = "\n".join(relevant_chunks)  # Combine chunks into context
        response = generate_response(query, context, tokenizer, model)  # Generate response
        st.write("Response:", response)  # Display response
    except Exception as e:
        st.error(f"An error occurred: {e}")  # Handle errors



